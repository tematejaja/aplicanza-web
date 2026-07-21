from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "entregables" / "Manual_de_Marca_Aplicanza.docx"
ASSETS = ROOT / "entregables" / "manual-marca-assets"
PORTRAIT = ROOT / "public" / "images" / "nicolas-alvarez.png"

NAVY = "0B1F3A"
TEAL = "138A8A"
BLUE = "2474E5"
PAPER = "F4F7F8"
SURFACE = "FDFEFE"
MUTED = "5D6B7C"
LINE = "D7E1E5"
WHITE = "FFFFFF"
RED = "A43D46"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.35)

def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def margins(cell, top=120, start=150, bottom=120, end=150):
    tc = cell._tc.get_or_add_tcPr()
    tc_mar = tc.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def borders(cell, color=LINE, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = tc_borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            tc_borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)

def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")

def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(index, len(widths) - 1)])
            margins(cell)

def set_font(run, name="Manrope", size=10, color=NAVY, bold=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.bold = bold

def add_text(text, size=10, color=NAVY, bold=False, font="Manrope"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.16
    set_font(p.add_run(text), font, size, color, bold)
    return p

def add_kicker(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text.upper())
    set_font(run, "IBM Plex Mono", 8, TEAL, True)
    run.font.letter_spacing = Pt(0.7)
    return p

def add_title(text, subtitle=None):
    p = doc.add_paragraph()
    p.style = styles["Heading 1"]
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.keep_with_next = True
    set_font(p.add_run(text), "Manrope", 27, NAVY, True)
    if subtitle:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(16)
        set_font(sp.add_run(subtitle), "Manrope", 11.5, MUTED, False)
    return p

def add_heading(text, level=2):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p

def add_rule(color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def add_callout(label, text, fill=PAPER, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9936])
    cell = table.cell(0, 0)
    shade(cell, fill)
    borders(cell, accent, 8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(label.upper() + "  "), "IBM Plex Mono", 7.5, accent, True)
    set_font(p.add_run(text), "Manrope", 11.5, NAVY, True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_two_col(cards):
    table = doc.add_table(rows=(len(cards) + 1) // 2, cols=2)
    table_geometry(table, [4888, 4888])
    index = 0
    for row in table.rows:
        for cell in row.cells:
            shade(cell, SURFACE)
            borders(cell, LINE, 6)
            if index < len(cards):
                title, body = cards[index]
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(4)
                set_font(p.add_run(title), "Manrope", 11, NAVY, True)
                p2 = cell.add_paragraph()
                p2.paragraph_format.space_after = Pt(0)
                set_font(p2.add_run(body), "Manrope", 9, MUTED, False)
            index += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def add_page(kicker, title, subtitle=None):
    if len(doc.paragraphs) > 0:
        doc.add_page_break()
    add_kicker(kicker)
    add_title(title, subtitle)
    add_rule()

def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(p.add_run("APLICANZA  ·  MANUAL DE MARCA  ·  "), "IBM Plex Mono", 7, MUTED, True)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Manrope"
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor.from_string(NAVY)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.16
for name, size, before, after in (("Heading 1", 20, 16, 8), ("Heading 2", 14, 12, 6), ("Heading 3", 11, 8, 4)):
    style = styles[name]
    style.font.name = "Manrope"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(NAVY if name != "Heading 3" else TEAL)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

add_footer(section)

# 1. Cover
doc.add_paragraph().paragraph_format.space_after = Pt(50)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(ASSETS / "aplicanza-logo.png"), width=Inches(4.9))
doc.add_paragraph().paragraph_format.space_after = Pt(35)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("MANUAL DE MARCA"), "IBM Plex Mono", 11, TEAL, True)
p.paragraph_format.space_after = Pt(12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Identidad visual, verbal y digital"), "Manrope", 25, NAVY, True)
p.paragraph_format.space_after = Pt(18)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("De la posibilidad a la implementación."), "Manrope", 13, MUTED, False)
p.paragraph_format.space_after = Pt(150)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("VERSIÓN 1.0  ·  JULIO DE 2026"), "IBM Plex Mono", 8, MUTED, True)

# 2. Essence
add_page("01 · Fundamento", "La marca nace para hacer avanzar.", "Aplicanza convierte la inteligencia artificial en mejoras concretas, medibles y adoptables para pymes.")
add_callout("Propósito", "Cerrar la distancia entre la posibilidad tecnológica y una implementación que produzca valor real.")
add_two_col([
    ("Promesa", "Comprender el proceso, probar en pequeño y medir antes de ampliar."),
    ("Posicionamiento", "Consultora cercana de implementación de IA para pymes, con criterio de negocio, datos y control de riesgos."),
    ("Problema que enfrentamos", "Empresas que pierden tiempo, dinero o información y no saben dónde aplicar IA con sentido."),
    ("Transformación", "Procesos más rentables, decisiones mejor informadas y equipos capaces de adaptarse."),
])
add_heading("Público prioritario", 2)
add_text("Gerentes, fundadores y líderes comerciales, operativos o administrativos de pymes. No necesitan conocer de inteligencia artificial; sí reconocer un proceso que puede funcionar mejor.")
add_heading("Territorio de marca", 2)
add_text("Aplicanza no vende futurismo ni herramientas aisladas. Habla de procesos, evidencia, personas, control y resultados.")

# 3. Message architecture
add_page("02 · Mensajes", "Tres frases, tres funciones.", "La consistencia verbal empieza por saber cuándo usar cada expresión de la marca.")
add_callout("Promesa principal", "Implementamos inteligencia artificial para que tu empresa avance.", PAPER, TEAL)
add_callout("Descriptor del logotipo", "IA aplicada para pymes.", "E9F5F4", TEAL)
add_callout("Lema comercial", "De la posibilidad a la implementación.", "EDF3FD", BLUE)
add_heading("Propuesta de valor", 2)
add_text("Encontramos dónde una empresa está perdiendo tiempo, dinero o información; usamos sus datos y conocimiento para diseñar una solución; la probamos en pequeño y demostramos el resultado antes de ampliarla.", 11)
add_heading("Jerarquía recomendada", 2)
add_two_col([
    ("Portada y campañas", "Promesa principal + llamada a solicitar diagnóstico."),
    ("Firma corporativa", "Logotipo + descriptor breve."),
    ("Presentaciones comerciales", "Lema comercial como cierre o transición."),
    ("Explicación extensa", "Propuesta de valor y metodología en lenguaje simple."),
])

# 4. Personality
add_page("03 · Personalidad", "Ejecutiva, tecnológica y cercana.", "La sofisticación de Aplicanza proviene del criterio y la claridad, no del exceso visual.")
add_two_col([
    ("Concreta", "Habla de procesos, tiempos, costos, datos e indicadores."),
    ("Cercana", "Acompaña a las personas que viven el proceso y explica sin distancia técnica."),
    ("Rigurosa", "Distingue evidencia, supuestos, riesgos y resultados verificables."),
    ("Adaptable", "Integra herramientas comerciales y open source según el contexto."),
    ("Serena", "Evita promesas grandilocuentes, urgencia artificial y espectáculo tecnológico."),
    ("Responsable", "Mantiene supervisión humana, límites y trazabilidad."),
])
add_heading("La tensión que define la marca", 2)
add_callout("Equilibrio", "Modernidad tecnológica + criterio ejecutivo + acompañamiento humano.")
add_heading("No somos", 2)
add_text("Una plantilla SaaS, una fábrica de chatbots, una promesa de automatización total ni una consultoría abstracta que termina en diapositivas.")

# 5. Logo
add_page("04 · Logotipo", "Un avance que atraviesa la A.", "El símbolo une origen, dirección y movimiento: la tecnología tiene sentido cuando hace avanzar el proceso.")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(ASSETS / "aplicanza-logo.png"), width=Inches(5.5))
p.paragraph_format.space_after = Pt(20)
add_heading("Versiones oficiales", 2)
table = doc.add_table(rows=1, cols=2)
table_geometry(table, [4888, 4888])
for cell in table.rows[0].cells:
    shade(cell, SURFACE); borders(cell)
p = table.cell(0, 0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(ASSETS / "aplicanza-mark.png"), width=Inches(1.45))
p = table.cell(0, 0).add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Símbolo"), "IBM Plex Mono", 8, TEAL, True)
p = table.cell(0, 1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(ASSETS / "aplicanza-logo-descriptor.png"), width=Inches(2.95))
p = table.cell(0, 1).add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Firma con descriptor"), "IBM Plex Mono", 8, TEAL, True)
add_heading("Área de protección", 2)
add_text("Conserve alrededor del logotipo un espacio mínimo equivalente a la altura de la barra horizontal de la flecha. Ningún texto, borde o imagen debe entrar en esa zona.")
add_heading("Tamaño mínimo", 2)
add_text("Firma principal: 120 px en digital o 32 mm en impresión. Símbolo: 24 px en digital o 8 mm en impresión. Debajo de esos tamaños use una versión simplificada y verifique legibilidad.")

# 6. Logo usage
add_page("05 · Uso del logotipo", "Reconocible en cualquier contexto.", "La marca debe conservar proporción, contraste y respiración.")
add_two_col([
    ("Sí · Fondo claro", "Use la firma oficial en azul profundo con el acento verde azulado."),
    ("Sí · Fondo oscuro", "Use una versión monocromática blanca o clara, preservando el símbolo."),
    ("Sí · Espacio suficiente", "Respete el área de protección y alinee ópticamente, no por intuición."),
    ("Sí · Archivos originales", "Utilice SVG en web y PNG de alta resolución en documentos."),
])
add_heading("Usos incorrectos", 2)
add_two_col([
    ("No deformar", "No estire, comprima, rote ni altere las proporciones."),
    ("No recolorear", "No introduzca degradados, neón, morado o colores fuera de la paleta."),
    ("No decorar", "No agregue sombras, biseles, contornos ni efectos tridimensionales."),
    ("No encerrar", "No coloque el logotipo dentro de cápsulas o recuadros innecesarios."),
    ("No competir", "Evite fondos fotográficos complejos o con poco contraste."),
    ("No reescribir", "No sustituya el nombre por otra tipografía ni modifique el símbolo."),
])

# 7. Color
add_page("06 · Color", "Confianza, claridad y avance.", "La paleta evita el cliché púrpura de la inteligencia artificial y construye una presencia sobria y reconocible.")
colors = [
    ("Azul profundo", NAVY, "#0B1F3A", "Texto, fondos oscuros, autoridad"),
    ("Verde Aplicanza", TEAL, "#138A8A", "Acciones, progreso, énfasis"),
    ("Azul evidencia", BLUE, "#2474E5", "Datos, enlaces, señal secundaria"),
    ("Papel", PAPER, "#F4F7F8", "Fondo principal"),
    ("Superficie", SURFACE, "#FDFEFE", "Tarjetas y documentos"),
    ("Texto secundario", MUTED, "#5D6B7C", "Descripciones y metadatos"),
]
table = doc.add_table(rows=1, cols=3)
table_geometry(table, [2100, 2500, 5336])
headers = ["Muestra", "Color", "Uso"]
for i, h in enumerate(headers):
    shade(table.rows[0].cells[i], NAVY); borders(table.rows[0].cells[i], NAVY)
    set_font(table.rows[0].cells[i].paragraphs[0].add_run(h), "Manrope", 9, WHITE, True)
set_repeat_table_header(table.rows[0])
for name, code, hexcode, use in colors:
    cells = table.add_row().cells
    for i, cell in enumerate(cells): borders(cell); margins(cell)
    shade(cells[0], code)
    set_font(cells[1].paragraphs[0].add_run(name), "Manrope", 9, NAVY, True)
    p = cells[1].add_paragraph(); set_font(p.add_run(hexcode), "IBM Plex Mono", 8, MUTED)
    set_font(cells[2].paragraphs[0].add_run(use), "Manrope", 9, MUTED)
table_geometry(table, [2100, 2500, 5336])
add_heading("Proporción orientativa", 2)
add_text("60% fondos claros · 25% azul profundo · 10% verde Aplicanza · 5% azul evidencia. El acento debe guiar, no dominar.")
add_callout("Accesibilidad", "Verifique contraste WCAG AA: 4,5:1 para texto normal y 3:1 para texto grande o elementos esenciales.")

# 8. Typography
add_page("07 · Tipografía", "Una voz clara con precisión técnica.", "Manrope sostiene la lectura; IBM Plex Mono identifica método, datos y metainformación.")
add_heading("Manrope · Tipografía principal", 2)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
set_font(p.add_run("Implementamos inteligencia artificial"), "Manrope", 23, NAVY, True)
add_text("Títulos: 600–700 · Texto: 400–500 · Botones y etiquetas: 650–750. Use interlineado compacto en titulares y cómodo en párrafos.")
add_heading("IBM Plex Mono · Tipografía de apoyo", 2)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
set_font(p.add_run("DOUBLE DIAMOND · LEAN · CRISP-DM · AGILE · NIST"), "IBM Plex Mono", 10, TEAL, True)
add_text("Úsela en etiquetas, pasos, cifras, coordenadas, fechas y marcos metodológicos. No la use para párrafos largos.")
add_heading("Escala digital recomendada", 2)
add_two_col([
    ("H1 · 56–80 px", "Promesa principal; una sola vez por página."),
    ("H2 · 38–64 px", "Inicio de secciones estratégicas."),
    ("H3 · 20–24 px", "Servicios, beneficios y componentes."),
    ("Cuerpo · 16–19 px", "Lectura clara con 1,5–1,65 de interlineado."),
])
add_callout("Regla", "La jerarquía depende del contraste de tamaño, peso y espacio; no de usar muchos colores.")

# 9. Imagery
add_page("08 · Fotografía e imagen", "Personas reales, trabajo real.", "La imagen debe comunicar acompañamiento, criterio y contexto; nunca espectáculo tecnológico vacío.")
table = doc.add_table(rows=1, cols=2)
table_geometry(table, [3500, 6276])
for cell in table.rows[0].cells: shade(cell, SURFACE); borders(cell)
p = table.cell(0, 0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(str(PORTRAIT), width=Inches(2.05))
p = table.cell(0, 1).paragraphs[0]
set_font(p.add_run("Dirección fotográfica"), "Manrope", 13, NAVY, True)
for title, body in [
    ("Natural", "Luz suave, textura real, expresiones auténticas."),
    ("Editorial", "Composición limpia, fondos sobrios y espacio negativo."),
    ("Contextual", "Personas revisando procesos, información o decisiones."),
    ("Diversa", "Representación verosímil de pymes y equipos latinoamericanos."),
]:
    q = table.cell(0, 1).add_paragraph(); q.paragraph_format.space_after = Pt(5)
    set_font(q.add_run(title + ". "), "Manrope", 9, TEAL, True)
    set_font(q.add_run(body), "Manrope", 9, MUTED)
add_heading("Evitar", 2)
add_text("Robots humanoides, cerebros luminosos, circuitos flotantes, manos estrechando en oficinas genéricas, pantallas imposibles, piel plástica, sonrisas excesivas o imágenes que aparenten una escala empresarial inexistente.")
add_heading("Capturas de producto", 2)
add_text("Muestre interfaces reales, legibles y con contexto. Incluya una breve explicación de qué decisión permite tomar la evidencia mostrada.")

# 10. Icons and graphics
add_page("09 · Lenguaje gráfico", "La tecnología se representa como sistema.", "Líneas, nodos y recorridos explican mejor el trabajo de Aplicanza que los efectos futuristas.")
add_two_col([
    ("Iconografía", "Trazo simple o duotono, geometría coherente y peso visual uniforme."),
    ("Diagramas", "Procesos de izquierda a derecha, puntos de control visibles y resultados al final."),
    ("Datos", "Escalas honestas, unidades explícitas, fuentes identificadas y metodología disponible."),
    ("Textura", "Grano muy sutil y cuadrículas técnicas solo como apoyo ambiental."),
    ("Bordes", "1 px, tono azul grisáceo; separan sin convertir cada bloque en una caja."),
    ("Movimiento", "Sutil, funcional y breve; debe mostrar avance, relación o cambio de estado."),
])
add_heading("Sistema de formas", 2)
add_text("Predominan rectángulos de radio moderado (8–12 px), líneas rectas y círculos pequeños como nodos. Evite cápsulas omnipresentes, esquinas excesivamente redondas y sombras profundas.")
add_callout("Principio", "Cada elemento gráfico debe explicar una relación, facilitar una acción o reforzar la jerarquía.")

# 11. Tone
add_page("10 · Voz y tono", "Hablar como un socio que entiende el negocio.", "La voz es clara, específica, respetuosa y orientada a decisiones.")
table = doc.add_table(rows=1, cols=2)
table_geometry(table, [4888, 4888])
for i, h in enumerate(("Preferir", "Evitar")):
    shade(table.rows[0].cells[i], TEAL if i == 0 else NAVY)
    borders(table.rows[0].cells[i], TEAL if i == 0 else NAVY)
    set_font(table.rows[0].cells[i].paragraphs[0].add_run(h), "Manrope", 10, WHITE, True)
set_repeat_table_header(table.rows[0])
pairs = [
    ("Reducimos tareas repetitivas y medimos el tiempo recuperado.", "Revolucionamos tu negocio con IA de última generación."),
    ("Probamos un proceso acotado antes de ampliar.", "Automatizamos todo tu negocio en semanas."),
    ("La supervisión humana permanece en decisiones sensibles.", "Nuestra IA elimina los errores humanos."),
    ("Revisamos si los datos son suficientes para el objetivo.", "Tus datos se convierten automáticamente en valor."),
]
for good, bad in pairs:
    cells = table.add_row().cells
    for c in cells: borders(c); margins(c)
    set_font(cells[0].paragraphs[0].add_run(good), "Manrope", 9, NAVY)
    set_font(cells[1].paragraphs[0].add_run(bad), "Manrope", 9, MUTED)
table_geometry(table, [4888, 4888])
add_heading("Reglas de redacción", 2)
add_text("Empiece por el problema o el resultado. Use verbos concretos. Explique las siglas la primera vez. Diferencie hechos, estimaciones y aspiraciones. Termine con un siguiente paso claro.")
add_callout("Llamada a la acción principal", "Solicitar diagnóstico · Continuar por WhatsApp.")

# 12. Methodology
add_page("11 · Metodología", "Primero comprendemos. Después ampliamos.", "Seis marcos conectados en una sola ruta de implementación.")
steps = [
    ("01 · Comprender", "Double Diamond", "Observamos el proceso con quienes lo conocen y definimos el problema antes de hablar de herramientas."),
    ("02 · Priorizar", "Lean", "Ordenamos oportunidades por impacto, esfuerzo, información disponible y riesgo."),
    ("03 · Trabajar los datos", "CRISP-DM", "Preparamos y validamos la información que sostiene la solución y sus decisiones."),
    ("04 · Pilotear", "Agile", "Construimos una versión acotada, la probamos en el trabajo real y aprendemos rápido."),
    ("05 · Controlar", "NIST AI RMF", "Definimos responsables, supervisión, límites y evidencias para operar con confianza."),
    ("06 · Adoptar", "Gestión del cambio", "Acompañamos a las personas para convertir la solución en una capacidad de la empresa."),
]
table = doc.add_table(rows=1, cols=3)
table_geometry(table, [2300, 2100, 5336])
for i, h in enumerate(("Etapa", "Marco", "Aplicación")):
    shade(table.rows[0].cells[i], NAVY); borders(table.rows[0].cells[i], NAVY)
    set_font(table.rows[0].cells[i].paragraphs[0].add_run(h), "Manrope", 9, WHITE, True)
set_repeat_table_header(table.rows[0])
for stage, framework, copy in steps:
    cells = table.add_row().cells
    for c in cells: borders(c); margins(c)
    set_font(cells[0].paragraphs[0].add_run(stage), "Manrope", 8.7, NAVY, True)
    set_font(cells[1].paragraphs[0].add_run(framework), "IBM Plex Mono", 7.8, TEAL, True)
    set_font(cells[2].paragraphs[0].add_run(copy), "Manrope", 8.5, MUTED)
table_geometry(table, [2300, 2100, 5336])
add_callout("Relato corto", "Comprendemos, priorizamos, trabajamos los datos, piloteamos, controlamos y acompañamos la adopción.")

# 13. Digital
add_page("12 · Experiencia digital", "Una interfaz que demuestra el método.", "La web debe sentirse ejecutiva y contemporánea, sin parecer una plantilla SaaS genérica.")
add_two_col([
    ("Composición", "Retículas editoriales, espacios amplios y jerarquía asimétrica con intención."),
    ("Componentes", "Tokens propios; tarjetas solo cuando ayudan a comparar o agrupar."),
    ("Interacción", "Estados claros, foco visible, navegación por teclado y respuesta inmediata."),
    ("Animación", "Revela recorridos, prioridades o causalidad; respeta movimiento reducido."),
    ("Contenido", "Resultados, metodología, casos verificables y responsable visible."),
    ("Conversión", "Formulario breve, WhatsApp y expectativa clara sobre el siguiente paso."),
])
add_heading("Principios técnicos", 2)
add_text("Responsive desde 320 px · contraste WCAG AA · HTML semántico · imágenes optimizadas · metadatos bilingües · datos estructurados · Core Web Vitals como criterio de publicación.")
add_heading("Llamadas a la acción", 2)
add_callout("Principal", "Solicitar diagnóstico", "E9F5F4", TEAL)
add_callout("Secundaria", "Ver casos · Conocer la metodología · Hablar por WhatsApp", "EDF3FD", BLUE)

# 14. Governance
add_page("13 · Gobernanza", "Consistencia antes de publicar.", "Este manual es la referencia para web, propuestas, presentaciones, redes y documentos comerciales.")
add_heading("Lista de control", 2)
checks = [
    "El logotipo proviene de un archivo oficial y conserva su proporción.",
    "Los colores pertenecen a la paleta y tienen contraste suficiente.",
    "Manrope e IBM Plex Mono cumplen su función asignada.",
    "El mensaje empieza por un problema, resultado o decisión concreta.",
    "Las afirmaciones sobre resultados tienen evidencia o se presentan como objetivos.",
    "La imagen muestra personas, procesos o productos reales sin clichés de IA.",
    "La pieza incluye un siguiente paso claro y un responsable identificable.",
    "La versión final fue revisada en móvil, escritorio e impresión cuando corresponda.",
]
table = doc.add_table(rows=0, cols=2)
table_geometry(table, [700, 9076])
for check in checks:
    cells = table.add_row().cells
    for c in cells: borders(c); margins(c)
    shade(cells[0], "E9F5F4")
    p = cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("✓"), "Manrope", 11, TEAL, True)
    set_font(cells[1].paragraphs[0].add_run(check), "Manrope", 9.2, NAVY)
table_geometry(table, [700, 9076])
add_heading("Responsabilidad", 2)
add_text("Toda nueva aplicación importante debe ser revisada por la dirección de Aplicanza. Si una necesidad no está cubierta, se documenta la decisión y se actualiza el manual en una nueva versión.")
add_callout("Contacto de marca", "Nicolás Álvarez · nicolasalvarezbernal@gmail.com · +57 300 296 8009")
add_text("Aplicanza · IA aplicada para pymes.\nBogotá · acompañamiento virtual en otras ciudades.", 8.5, MUTED, False, "IBM Plex Mono")

doc.core_properties.title = "Manual de Marca Aplicanza"
doc.core_properties.subject = "Identidad visual, verbal y digital"
doc.core_properties.author = "Aplicanza"
doc.core_properties.keywords = "Aplicanza, marca, identidad, inteligencia artificial, pymes"
image_alts = [
    "Logotipo principal de Aplicanza",
    "Logotipo principal de Aplicanza en la sección de identidad",
    "Símbolo de Aplicanza",
    "Logotipo de Aplicanza con el descriptor IA aplicada para pymes",
    "Retrato de Nicolás Álvarez, fundador de Aplicanza",
]
for shape, alt in zip(doc.inline_shapes, image_alts):
    shape._inline.docPr.set("descr", alt)
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
